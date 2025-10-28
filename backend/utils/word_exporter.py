from docx import Document
from pathlib import Path

def save_as_word(title: str, content: str, filename: str, out_dir: str):
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(content)

    out_dir_path = Path(out_dir)
    out_dir_path.mkdir(parents=True, exist_ok=True)
    path = out_dir_path / f"{filename}.docx"
    doc.save(str(path))
    return str(path)
